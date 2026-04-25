"""
Tests pour le module Extractor.
Couvre les services d'extraction (PDF, DOCX), la normalisation des labels,
et les endpoints API.

Lancer avec : pytest tests/ -v
"""
import pytest
import os
import json
import tempfile
from docx import Document as DocxDocument
from docx.oxml.ns import qn

# ─── Fixtures ────────────────────────────────────────────


@pytest.fixture
def app():
    """Crée une instance de l'application Flask en mode test."""
    # On doit être dans le bon répertoire pour les imports
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
    from app import create_app
    app = create_app()
    app.config['TESTING'] = True
    yield app


@pytest.fixture
def client(app):
    """Client HTTP Flask pour les tests d'API."""
    return app.test_client()


@pytest.fixture
def sample_docx(tmp_path):
    """Crée un fichier DOCX de test avec du texte et un tableau."""
    doc = DocxDocument()
    doc.add_paragraph("Titre du document")
    doc.add_paragraph("Paragraphe de texte normal.")

    # Ajouter un tableau 3x2
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Nom"
    table.cell(0, 1).text = "Valeur"
    table.cell(1, 0).text = "Position"
    table.cell(1, 1).text = "8701.20.00.10"
    table.cell(2, 0).text = "Désignation"
    table.cell(2, 1).text = "Tracteurs agricoles"

    doc.add_paragraph("Paragraphe final.")

    filepath = str(tmp_path / "test_document.docx")
    doc.save(filepath)
    return filepath


@pytest.fixture
def sample_docx_with_pages(tmp_path):
    """Crée un DOCX avec un saut de page explicite."""
    doc = DocxDocument()
    doc.add_paragraph("Page 1 - Premier paragraphe")

    # Ajouter un saut de page
    p = doc.add_paragraph()
    run = p.add_run()
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)

    doc.add_paragraph("Page 2 - Après saut de page")

    filepath = str(tmp_path / "test_multipage.docx")
    doc.save(filepath)
    return filepath


@pytest.fixture
def sample_docx_with_tabs(tmp_path):
    """Crée un DOCX avec des pseudo-tableaux tabulés."""
    doc = DocxDocument()
    doc.add_paragraph("En-tête")
    doc.add_paragraph("Col1\tCol2\tCol3")
    doc.add_paragraph("A\tB\tC")
    doc.add_paragraph("D\tE\tF")
    doc.add_paragraph("Pied de page")

    filepath = str(tmp_path / "test_tabs.docx")
    doc.save(filepath)
    return filepath


# ═══════════════════════════════════════════════════════════
# Tests : DOCX Extractor
# ═══════════════════════════════════════════════════════════

class TestDocxExtractor:
    """Tests pour le service d'extraction DOCX."""

    def test_extract_basic_document(self, sample_docx):
        """Vérifie l'extraction d'un DOCX avec texte et tableau."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx)

        assert len(content) >= 3  # au moins 2 textes + 1 tableau
        types = [b['type'] for b in content]
        assert 'texte' in types
        assert 'tableau' in types

    def test_text_blocks_have_content(self, sample_docx):
        """Vérifie que les blocs texte ont un contenu."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx)
        textes = [b for b in content if b['type'] == 'texte']

        for t in textes:
            assert 'contenu' in t
            assert len(t['contenu']) > 0

    def test_table_blocks_have_rows(self, sample_docx):
        """Vérifie que les blocs tableau ont des lignes."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx)
        tableaux = [b for b in content if b['type'] == 'tableau']

        assert len(tableaux) >= 1
        for tab in tableaux:
            assert 'lignes' in tab
            assert len(tab['lignes']) > 0
            assert 'numero' in tab

    def test_blocks_have_page_number(self, sample_docx):
        """Vérifie que tous les blocs ont un numéro de page."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx)

        for block in content:
            assert 'page' in block, f"Bloc sans page: {block}"
            assert block['page'] >= 1

    def test_page_break_detection(self, sample_docx_with_pages):
        """Vérifie que les sauts de page Word sont détectés."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx_with_pages)

        pages = set(b['page'] for b in content)
        assert len(pages) >= 2, f"Attendu >= 2 pages, obtenu: {pages}"

    def test_tab_pseudo_table_detection(self, sample_docx_with_tabs):
        """Vérifie la détection des pseudo-tableaux tabulés."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx_with_tabs)

        tableaux = [b for b in content if b['type'] == 'tableau']
        assert len(tableaux) >= 1, "Aucun pseudo-tableau détecté"

    def test_column_filtering(self, sample_docx):
        """Vérifie le filtrage par colonnes."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx, table_columns=[0])
        tableaux = [b for b in content if b['type'] == 'tableau']

        for tab in tableaux:
            for row in tab['lignes']:
                assert 'col_00' in row
                assert 'col_01' not in row

    def test_extraction_order_preserved(self, sample_docx):
        """Vérifie que l'ordre texte/tableau est préservé."""
        from app.services.docx_extractor import extract_document
        content = extract_document(sample_docx)

        # Le premier bloc doit être du texte (titre)
        assert content[0]['type'] == 'texte'
        assert 'Titre' in content[0]['contenu']


# ═══════════════════════════════════════════════════════════
# Tests : Normalisation des labels
# ═══════════════════════════════════════════════════════════

class TestNormalizeLabels:
    """Tests pour la normalisation des étiquettes."""

    def test_basic_renaming(self):
        """Vérifie le renommage basique des clés."""
        from app.services.pdf_extractor import normalize_labels
        data = [{"col_04": "Tracteurs", "col_05": "D.D 30%"}]
        result = normalize_labels(data)

        assert len(result) == 1
        row = result[0]
        assert "Désignation" in row
        assert row["Désignation"] == "Tracteurs"

    def test_value_cleaning(self):
        """Vérifie le nettoyage des valeurs (suppression des préfixes)."""
        from app.services.pdf_extractor import normalize_labels
        data = [{"col_05": "D.D 30%"}]
        result = normalize_labels(data)

        row = result[0]
        assert "DD" in row
        assert "D.D" not in row["DD"]

    def test_position_cleaning(self):
        """Vérifie le nettoyage de la colonne Position."""
        from app.services.pdf_extractor import normalize_labels
        data = [{"Position & Sous Position": "Position 8701.20.00"}]
        result = normalize_labels(data)

        assert result[0]["Position"] == "8701.20.00"

    def test_collision_handling(self):
        """Vérifie la gestion des collisions de clés."""
        from app.services.pdf_extractor import normalize_labels
        data = [{"col_08": "valeur1", "F.A.P": "valeur2"}]
        result = normalize_labels(data)

        row = result[0]
        assert "FAP" in row
        # Les deux colonnes mappent vers FAP, donc concaténation
        assert "valeur1" in row["FAP"] or "valeur2" in row["FAP"]

    def test_unknown_keys_preserved(self):
        """Vérifie que les clés inconnues sont conservées."""
        from app.services.pdf_extractor import normalize_labels
        data = [{"custom_field": "ma_valeur"}]
        result = normalize_labels(data)

        assert result[0]["custom_field"] == "ma_valeur"

    def test_empty_input(self):
        """Vérifie le comportement avec une liste vide."""
        from app.services.pdf_extractor import normalize_labels
        result = normalize_labels([])
        assert result == []

    def test_custom_mapping_name(self):
        """Vérifie le chargement du mapping par défaut (fallback)."""
        from app.services.pdf_extractor import normalize_labels
        # Un mapping inexistant doit utiliser le fallback sans erreur
        data = [{"col_04": "test"}]
        result = normalize_labels(data, mapping_name="inexistant")
        assert len(result) == 1


# ═══════════════════════════════════════════════════════════
# Tests : Extraction de codes tarifaires
# ═══════════════════════════════════════════════════════════

class TestTariffExtraction:
    """Tests pour l'extraction dynamique de codes tarifaires."""

    def test_tariff_pattern_detection(self):
        """Vérifie la détection du format XXXX.XX.XX.XX."""
        from app.services.pdf_extractor import extract_rows_with_single_tariff_code
        # La fonction filtre sur les tableaux contenant "col_04" ou "désignation des produits"
        content = [
            {
                "type": "tableau",
                "lignes": [
                    {"col_00": "8701.20.00.10", "col_04": "Tracteurs"},
                    {"col_00": "Texte normal", "col_04": "Sans code"},
                    {"col_00": "8702.10.00.20", "col_04": "Véhicules"},
                ]
            }
        ]
        result = extract_rows_with_single_tariff_code(content)
        # Les 3 lignes sont valides : 2 avec un code unique, 1 sans code (codes_found=0 <= 1)
        assert len(result) == 3

    def test_no_tariff_codes(self):
        """Vérifie le comportement sans codes tarifaires."""
        from app.services.pdf_extractor import extract_rows_with_single_tariff_code
        content = [
            {
                "type": "tableau",
                "lignes": [
                    {"col_00": "Pas de code", "col_01": "Normal"}
                ]
            }
        ]
        result = extract_rows_with_single_tariff_code(content)
        assert len(result) == 0


# ═══════════════════════════════════════════════════════════
# Tests : Config labels (chargement JSON)
# ═══════════════════════════════════════════════════════════

class TestLabelConfig:
    """Tests pour le chargement de la config de normalisation."""

    def test_default_config_loads(self):
        """Vérifie que le fichier default.json se charge correctement."""
        from app.services.pdf_extractor import _load_label_config
        mapping, cleaning = _load_label_config("default")

        assert isinstance(mapping, dict)
        assert isinstance(cleaning, dict)
        assert len(mapping) > 0
        assert "Position & Sous Position" in mapping or "col_04" in mapping

    def test_fallback_on_missing_file(self):
        """Vérifie le fallback quand le fichier n'existe pas."""
        from app.services.pdf_extractor import _load_label_config
        mapping, cleaning = _load_label_config("fichier_inexistant_xyz")

        assert isinstance(mapping, dict)
        assert len(mapping) > 0  # Les valeurs hardcodées de fallback

    def test_cleaning_config_structure(self):
        """Vérifie la structure de la config de nettoyage."""
        from app.services.pdf_extractor import _load_label_config
        _, cleaning = _load_label_config("default")

        for key, patterns in cleaning.items():
            assert isinstance(patterns, list), f"Les patterns de '{key}' doivent être une liste"
            for pattern in patterns:
                assert isinstance(pattern, str), f"Chaque pattern doit être un string"


# ═══════════════════════════════════════════════════════════
# Tests : API Endpoints
# ═══════════════════════════════════════════════════════════

class TestAPIEndpoints:
    """Tests d'intégration pour les endpoints Flask."""

    def test_index_endpoint(self, client):
        """Vérifie l'endpoint racine."""
        response = client.get('/')
        assert response.status_code == 200
        data = response.get_json()
        assert data['message'] == 'Document Extractor API is running'
        assert 'endpoints' in data

    def test_extract_document_no_file(self, client):
        """Vérifie l'erreur quand aucun fichier n'est fourni."""
        response = client.post('/api/extract-document')
        assert response.status_code == 400
        data = response.get_json()
        assert 'error' in data

    def test_extract_pdf_no_file(self, client):
        """Vérifie l'erreur quand aucun fichier n'est fourni pour le PDF."""
        response = client.post('/api/extract-pdf')
        assert response.status_code == 400

    def test_extract_docx_endpoint(self, client, sample_docx):
        """Vérifie l'extraction DOCX via l'API."""
        with open(sample_docx, 'rb') as f:
            response = client.post(
                '/api/extract-docx',
                data={'file': (f, 'test.docx')},
                content_type='multipart/form-data'
            )
        assert response.status_code == 200
        data = response.get_json()
        assert data['success'] is True
        assert 'contenu' in data
        assert data['total_blocs'] > 0

    def test_extract_document_docx(self, client, sample_docx):
        """Vérifie l'extraction unifiée avec un DOCX."""
        with open(sample_docx, 'rb') as f:
            response = client.post(
                '/api/extract-document',
                data={'file': (f, 'test.docx')},
                content_type='multipart/form-data'
            )
        assert response.status_code == 200
        data = response.get_json()
        assert data['success'] is True
        assert data['format'] == 'docx'

    def test_extract_document_unsupported_format(self, client, tmp_path):
        """Vérifie le rejet d'un format non supporté."""
        fake_file = tmp_path / "test.txt"
        fake_file.write_text("contenu texte")
        with open(str(fake_file), 'rb') as f:
            response = client.post(
                '/api/extract-document',
                data={'file': (f, 'test.txt')},
                content_type='multipart/form-data'
            )
        assert response.status_code == 400

    def test_normalize_labels_endpoint(self, client):
        """Vérifie l'endpoint de normalisation."""
        data = [{"col_04": "Tracteurs", "col_05": "D.D 30%"}]
        response = client.post(
            '/api/normalize-labels',
            data=json.dumps(data),
            content_type='application/json'
        )
        assert response.status_code == 200
        result = response.get_json()
        assert result['success'] is True
        assert result['mapping_utilise'] == 'default'
        assert len(result['donnees']) == 1

    def test_normalize_labels_custom_mapping(self, client):
        """Vérifie le paramètre mapping dans l'URL."""
        data = [{"col_04": "test"}]
        response = client.post(
            '/api/normalize-labels?mapping=custom',
            data=json.dumps(data),
            content_type='application/json'
        )
        assert response.status_code == 200
        result = response.get_json()
        assert result['mapping_utilise'] == 'custom'

    def test_normalize_labels_invalid_body(self, client):
        """Vérifie l'erreur avec un body invalide."""
        response = client.post(
            '/api/normalize-labels',
            data='not json',
            content_type='application/json'
        )
        assert response.status_code in [400, 500]
