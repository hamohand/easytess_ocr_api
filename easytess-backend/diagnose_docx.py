"""
Script de diagnostic pour analyser la structure d'un fichier Word.
Usage: python diagnose_docx.py chemin/vers/document.docx
"""
import sys
from docx import Document

def diagnose(docx_path):
    doc = Document(docx_path)
    
    print(f"=== Diagnostic de: {docx_path} ===\n")
    print(f"Nombre de paragraphes (doc.paragraphs): {len(doc.paragraphs)}")
    print(f"Nombre de tableaux (doc.tables): {len(doc.tables)}")
    
    # Compter les paragraphes avec tabulations
    tab_paragraphs = [p for p in doc.paragraphs if '\t' in p.text]
    print(f"Paragraphes contenant des tabulations: {len(tab_paragraphs)}")
    
    # Afficher les 20 premiers paragraphes
    print(f"\n--- 20 premiers paragraphes ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        text = p.text
        style_name = p.style.name if p.style else "(aucun style)"
        # Montrer les tabulations explicitement
        display = text.replace('\t', ' | ')
        has_tab = '[TAB]' if '\t' in text else ''
        tab_count = text.count('\t')
        print(f"  P{i:02d} ({tab_count} tabs): {display[:150]}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python diagnose_docx.py <fichier.docx>")
        sys.exit(1)
    diagnose(sys.argv[1])
