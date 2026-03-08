"""
Service de conversion PDF → Word (.docx)
Prend le contenu structuré extrait d'un PDF et le reconstruit
en document Word, en préservant la structure:
  - Paragraphes de texte
  - Tableaux avec en-têtes détectés
  - Sauts de page entre les pages PDF
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import logging
import os

logger = logging.getLogger(__name__)


def _add_table_style(table, has_header=False):
    """
    Applique un style visuel au tableau Word.
    """
    # Bordures sur toutes les cellules
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr %s/>' % nsdecls('w')
    )

    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

    # Style de l'en-tête si détecté
    if has_header and len(table.rows) > 0:
        for cell in table.rows[0].cells:
            # Fond gris clair pour l'en-tête
            shading = parse_xml(
                '<w:shd %s w:fill="E8E8E8" w:val="clear"/>' % nsdecls('w')
            )
            cell._tc.get_or_add_tcPr().append(shading)

            # Texte en gras
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)


def _add_page_break(doc):
    """Ajoute un saut de page."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx_break_type=None)
    # Saut de page via XML
    br = parse_xml('<w:br %s w:type="page"/>' % nsdecls('w'))
    run._r.append(br)


def convert_content_to_docx(content, output_path, source_filename=None):
    """
    Convertit le contenu structuré (texte + tableaux) en document Word.

    Args:
        content: Liste de blocs (depuis pdf_extractor ou docx_extractor)
                 - {"type": "texte", "contenu": "...", "page": N}
                 - {"type": "tableau", "numero": N, "page": N,
                    "lignes": [...], "metadata": {...}}
        output_path: Chemin de sortie pour le .docx
        source_filename: Nom du fichier source (optionnel, pour le titre)

    Returns:
        str: Chemin du fichier .docx créé
    """
    doc = Document()

    # ─── Configurer les styles par défaut ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ─── Titre optionnel ───
    if source_filename:
        heading = doc.add_heading(level=1)
        run = heading.add_run(f"Document converti : {source_filename}")
        run.font.size = Pt(14)
        doc.add_paragraph("")  # Espacement

    current_page = None

    for block in content:
        block_page = block.get("page")

        # Saut de page si on change de page PDF
        if block_page and current_page and block_page != current_page:
            _add_page_break(doc)
        current_page = block_page

        if block["type"] == "texte":
            _add_text_block(doc, block)

        elif block["type"] == "tableau":
            _add_table_block(doc, block)

    doc.save(output_path)
    logger.info(f"Document Word créé: {output_path}")
    return output_path


def _add_text_block(doc, block):
    """Ajoute un bloc de texte au document."""
    text = block.get("contenu", "")
    if not text:
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'


def _add_table_block(doc, block):
    """Ajoute un bloc tableau au document Word."""
    lignes = block.get("lignes", [])
    if not lignes:
        return

    metadata = block.get("metadata", {})
    has_header = metadata.get("a_entete", False)
    entetes = metadata.get("entetes", [])

    # Déterminer les colonnes à partir des données
    all_keys = []
    for ligne in lignes:
        for key in ligne.keys():
            if key not in all_keys:
                all_keys.append(key)

    if not all_keys:
        return

    nb_cols = len(all_keys)

    # Nombre de lignes = données + en-tête éventuel
    nb_rows = len(lignes) + (1 if entetes else 0)

    table = doc.add_table(rows=nb_rows, cols=nb_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_offset = 0

    # ─── En-têtes si disponibles ───
    if entetes:
        for col_idx, header_text in enumerate(entetes):
            if col_idx < nb_cols:
                cell = table.rows[0].cells[col_idx]
                cell.text = header_text
        row_offset = 1
    elif has_header:
        # Utiliser les noms de clés comme en-têtes
        for col_idx, key in enumerate(all_keys):
            cell = table.rows[0].cells[col_idx]
            # Formater la clé: "col_00" -> "Col 00", sinon garder tel quel
            display_key = key.replace('_', ' ').title()
            cell.text = display_key

    # ─── Données ───
    for row_idx, ligne in enumerate(lignes):
        actual_row = row_idx + row_offset
        if actual_row >= nb_rows:
            break
        for col_idx, key in enumerate(all_keys):
            if col_idx < nb_cols:
                cell = table.rows[actual_row].cells[col_idx]
                cell.text = str(ligne.get(key, ""))

    # ─── Style du tableau ───
    _add_table_style(table, has_header=bool(entetes) or has_header)

    # Espacement après le tableau
    doc.add_paragraph("")


def convert_pdf_to_docx(pdf_path, output_path=None):
    """
    Raccourci : extrait le contenu d'un PDF et le convertit directement en .docx.

    Args:
        pdf_path: Chemin vers le fichier .pdf
        output_path: Chemin de sortie .docx (optionnel, généré automatiquement)

    Returns:
        tuple: (chemin_docx, stats_extraction)
    """
    from app.services.pdf_extractor import extract_pdf

    # Générer le chemin de sortie si non fourni
    if not output_path:
        base, _ = os.path.splitext(pdf_path)
        output_path = f"{base}.docx"

    # Extraire le contenu du PDF
    content, stats = extract_pdf(pdf_path, strategy="auto", include_metadata=True)

    if not content:
        raise ValueError("Aucun contenu extractible trouvé dans le PDF")

    # Convertir en DOCX
    source_name = os.path.basename(pdf_path)
    convert_content_to_docx(content, output_path, source_filename=source_name)

    return output_path, stats
