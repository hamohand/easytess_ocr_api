"""
Script pour générer un document Word de test
contenant du texte et des tableaux.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Titre ---
doc.add_heading('Rapport de Test EasyTess', level=1)

# --- Paragraphe 1 ---
doc.add_paragraph(
    "Ce document est un exemple contenant du texte classique et des tableaux. "
    "Il est utilisé pour tester l'endpoint /api/extract-docx."
)

# --- Tableau 1 : Employés ---
doc.add_heading('Liste des employés', level=2)

table1 = doc.add_table(rows=4, cols=4, style='Table Grid')
# En-têtes
headers = ['Nom', 'Prénom', 'Ville', 'Poste']
for i, h in enumerate(headers):
    table1.rows[0].cells[i].text = h

# Données
data = [
    ['Dupont', 'Jean', 'Paris', 'Développeur'],
    ['Martin', 'Sophie', 'Lyon', 'Designer'],
    ['Bernard', 'Pierre', 'Marseille', 'Chef de projet'],
]
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, val in enumerate(row_data):
        table1.rows[row_idx].cells[col_idx].text = val

# --- Paragraphe entre les tableaux ---
doc.add_paragraph(
    "Le tableau ci-dessus présente les employés de l'entreprise. "
    "Le tableau suivant récapitule les ventes trimestrielles."
)

# --- Tableau 2 : Ventes ---
doc.add_heading('Ventes trimestrielles', level=2)

table2 = doc.add_table(rows=5, cols=3, style='Table Grid')
headers2 = ['Trimestre', 'Chiffre d\'affaires (€)', 'Objectif atteint']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h

data2 = [
    ['T1 2025', '125 000', 'Oui'],
    ['T2 2025', '98 500', 'Non'],
    ['T3 2025', '142 300', 'Oui'],
    ['T4 2025', '156 800', 'Oui'],
]
for row_idx, row_data in enumerate(data2, start=1):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx].cells[col_idx].text = val

# --- Conclusion ---
doc.add_paragraph(
    "Conclusion : les résultats de l'année 2025 sont globalement positifs."
)

# Sauvegarde
output_path = 'test_document.docx'
doc.save(output_path)
print(f"Document de test créé : {output_path}")
